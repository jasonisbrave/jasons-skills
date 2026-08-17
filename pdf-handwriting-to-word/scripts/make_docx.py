# -*- coding: utf-8 -*-
"""把转写好的纯文本（简单标记约定）转换为带中文公文格式的 Word 文档。

用法:
    python make_docx.py <输入.md/txt> <输出.docx>

输入文件约定（UTF-8，逐行解析，空行忽略）:
    # 标题        -> 居中、黑体、16pt、加粗（可用多行 # 作为主副标题）
    ## 节标题     -> 黑体、14pt、加粗、无缩进
    ### 条目标题  -> 仿宋、12pt、加粗、首行缩进2字符
    其他行        -> 正文：仿宋、12pt、首行缩进2字符、1.5倍行距

依赖:
    python-docx  ( pip install python-docx )
"""
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def add_para(doc, text, *, bold=False, center=False, size=12,
             indent=True, ea_font="仿宋", space_after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.5
    if indent and not center:
        pf.first_line_indent = Pt(size * 2)  # 首行缩进2字符
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea_font)
    return p


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    src, out = sys.argv[1], sys.argv[2]

    with open(src, encoding="utf-8") as f:
        lines = [ln.rstrip() for ln in f]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    for ln in lines:
        if not ln.strip():
            continue
        if ln.startswith("### "):
            add_para(doc, ln[4:].strip(), bold=True)
        elif ln.startswith("## "):
            add_para(doc, ln[3:].strip(), bold=True, size=14,
                     indent=False, ea_font="黑体", space_after=8)
        elif ln.startswith("# "):
            add_para(doc, ln[2:].strip(), bold=True, center=True, size=16,
                     indent=False, ea_font="黑体", space_after=12)
        else:
            add_para(doc, ln.strip())

    doc.save(out)
    print("saved:", out)


if __name__ == "__main__":
    main()
